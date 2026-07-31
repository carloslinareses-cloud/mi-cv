# -*- coding: utf-8 -*-
"""
Genera varias versiones del CV en formato ATS, una por tipo de candidatura.

    py generar_cv_ats.py

Por qué varias versiones y no una: los sistemas ATS puntúan por coincidencia de
palabras con la oferta. Un CV de desarrollador enviado a un puesto de asistente
virtual saca una puntuación baja aunque la persona valga, porque no comparte
vocabulario con la oferta.

Reglas de formato que NO hay que romper (son las que hacen que un ATS lo lea):
  - Una sola columna. Sin tablas, sin cuadros de texto, sin columnas.
  - Sin imágenes, sin iconos, sin gráficos de "nivel de habilidad".
  - Los datos de contacto en el cuerpo, nunca en el encabezado del documento:
    muchos ATS ni siquiera leen los encabezados.
  - Títulos de sección con los nombres de siempre. "OTRA EXPERIENCIA" NO está en
    la lista que reconocen los parsers: todo va en EXPERIENCIA PROFESIONAL.
  - Fechas en MM/AAAA. "Abr", "Ene" y "Dic" no existen en un léxico configurado
    en inglés, y ahí el parser se queda sin fecha de inicio o inventa un hueco.
  - Cabecera de cada puesto en DOS líneas, con el cargo solo en la primera. Si
    cargo, empresa y ciudad comparten línea y separador, el parser tiene que
    adivinar qué campo es cuál, y con un cargo que lleva comas lo parte en tres.

TODO el contenido sale de la experiencia real recogida en el CV publicado.
No se inventa ni un puesto, ni un año, ni una herramienta.
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SALIDA = Path(__file__).resolve().parent / "output" / "ats"
SALIDA.mkdir(parents=True, exist_ok=True)

NEGRO = RGBColor(0, 0, 0)
GRIS = RGBColor(70, 70, 70)

CONTACTO = {
    "nombre": "CARLOS HUMBERTO LINARES MORALES",
    "correo": "carlos.linares.es@gmail.com",
    "telefono": "+58 414 227 2792",
    # URL personalizada, confirmada por Carlos como la real y activa. Al
    # personalizarla en LinkedIn, la larga por defecto deja de funcionar.
    "linkedin": "linkedin.com/in/cahulin",
    "github": "github.com/carloslinareses-cloud",
    "web": "carloslinareses-cloud.github.io/mi-cv",
    "ubicacion": "Venezuela · Disponible para trabajo 100% remoto",
}


# --- utilidades de formato ---------------------------------------------------

def _fuente(run, tam=10.5, negrita=False, color=NEGRO, cursiva=False):
    run.font.name = "Arial"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Arial")
    rpr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(tam)
    run.bold = negrita
    run.italic = cursiva
    run.font.color.rgb = color


def _parrafo(doc, texto="", tam=10.5, negrita=False, color=NEGRO,
             antes=0, despues=2, cursiva=False, alineacion=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(despues)
    p.paragraph_format.line_spacing = 1.05
    if alineacion is not None:
        p.alignment = alineacion
    if texto:
        _fuente(p.add_run(texto), tam, negrita, color, cursiva)
    return p


def _seccion(doc, titulo):
    """Título de sección con una línea debajo, en mayúsculas y sin adornos."""
    p = _parrafo(doc, titulo, tam=11, negrita=True, antes=11, despues=3)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "808080")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def _punto(doc, texto):
    """Viñeta escrita a mano: las listas automáticas de Word a veces se pierden
    al convertir el documento, y el texto se junta en un solo bloque."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.05
    _fuente(p.add_run("• " + texto), 10.5)
    return p


def _cabecera(doc, mostrar_github=True):
    _parrafo(doc, CONTACTO["nombre"], tam=15, negrita=True, despues=2,
             alineacion=WD_ALIGN_PARAGRAPH.CENTER)
    enlaces = [CONTACTO["linkedin"]]
    if mostrar_github:
        enlaces.append(CONTACTO["github"])
    enlaces.append(CONTACTO["web"])
    for t in (f"{CONTACTO['correo']} | {CONTACTO['telefono']}",
              " | ".join(enlaces),
              CONTACTO["ubicacion"]):
        _parrafo(doc, t, tam=9.5, color=GRIS, despues=1,
                 alineacion=WD_ALIGN_PARAGRAPH.CENTER)


def _documento():
    doc = Document()
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Inches(0.55)
    s.left_margin = s.right_margin = Inches(0.65)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    return doc


def _metadatos(doc, especialidad):
    """Sin esto el documento sale firmado por 'python-docx' y con fecha de 2013,
    que es lo que ve cualquiera que abra Propiedades en Word."""
    p = doc.core_properties
    p.author = "Carlos Humberto Linares Morales"
    p.last_modified_by = "Carlos Humberto Linares Morales"
    p.title = f"CV Carlos Humberto Linares Morales - {especialidad}"
    p.comments = ""
    p.created = datetime.now()
    p.modified = datetime.now()


def construir(perfil):
    doc = _documento()
    _cabecera(doc, mostrar_github=perfil.get("github", True))

    _seccion(doc, "PERFIL PROFESIONAL")
    _parrafo(doc, perfil["titular"], tam=11, negrita=True, despues=3)
    _parrafo(doc, perfil["resumen"], despues=4)

    # Todos los puestos en una sola sección y en orden cronológico inverso.
    # Los menos relevantes llevan una viñeta en vez de tres, pero mantienen el
    # mismo formato: así el parser los reconoce como empleos y no se pierden.
    _seccion(doc, "EXPERIENCIA PROFESIONAL")
    breves = {e["puesto"] for e in perfil.get("breves", [])}
    todos = list(perfil["experiencia"]) + list(perfil.get("breves", []))
    for e in sorted(todos, key=lambda x: x["desde"], reverse=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        _fuente(p.add_run(e["puesto"]), 10.8, negrita=True)
        # Segunda línea: empresa | ubicación | fechas, con separador uniforme.
        _parrafo(doc, f"{e['empresa']} | {e['lugar']} | {e['fechas']}",
                 tam=9.5, color=GRIS, despues=2)
        logros = e["logros"][:1] if e["puesto"] in breves else e["logros"]
        for b in logros:
            _punto(doc, b)

    _seccion(doc, "HABILIDADES")
    for grupo, items in perfil["habilidades"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        _fuente(p.add_run(f"{grupo}: "), 10.5, negrita=True)
        _fuente(p.add_run(", ".join(items)), 10.5)

    if perfil.get("proyectos"):
        _seccion(doc, "PROYECTOS")
        for pr in perfil["proyectos"]:
            _punto(doc, pr)

    _seccion(doc, "IDIOMAS")
    _parrafo(doc, perfil["idiomas"])

    # Va ANTES de la formación: un encabezado que el parser no reconozca
    # colocado después de Educación hace que todo su contenido se indexe como
    # parte de la formación académica.
    _seccion(doc, "INFORMACIÓN ADICIONAL")
    for b in perfil["adicional"]:
        _punto(doc, b)

    _seccion(doc, "FORMACIÓN ACADÉMICA")
    _parrafo(doc, "TSU en Informática, IUT José María Carreño, 2011 – 2015")

    _metadatos(doc, perfil["especialidad"])
    ruta = SALIDA / perfil["archivo"]
    doc.save(ruta)
    return ruta


if __name__ == "__main__":
    from perfiles_ats import PERFILES

    print("Generando CV con formato ATS...\n")
    for p in PERFILES:
        ruta = construir(p)
        print(f"  OK  {ruta.name}")
    print(f"\nCarpeta: {SALIDA}")
